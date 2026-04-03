#!/usr/bin/env python3
"""Generate the Monday Night Soccer App – Workflow & Feature Requirements document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

# ── Title Page ──────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Monday Night Soccer App')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Workflow & Feature Requirements')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Version 1.0  •  {datetime.date.today().strftime("%d %B %Y")}').font.size = Pt(12)

doc.add_page_break()

# ── Table of Contents placeholder ───────────────────────────────────────────
doc.add_heading('Table of Contents', level=1)
doc.add_paragraph('(Update in Word: References → Update Table)', style='Normal')
# Insert TOC field
p = doc.add_paragraph()
run = p.add_run()
fld_char1 = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
run._r.append(fld_char1)
run2 = p.add_run(' TOC \\o "1-3" \\h \\z \\u ')
fld_char2 = run2._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
run2._r.append(fld_char2)
run3 = p.add_run('[Right-click → Update Field to populate]')
fld_char3 = run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
run3._r.append(fld_char3)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 1 – OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'This document defines the comprehensive feature set for the Monday Night Soccer session '
    'management application. It covers every stage of the session workflow, the actions available '
    'to admins and players at each stage, backward transitions between stages, and recommended '
    'improvements for flexibility and usability.'
)

doc.add_heading('1.1 Session Workflow Stages', level=2)
doc.add_paragraph(
    'Each session moves through the following lifecycle stages. '
    'Forward transitions are the primary flow; backward transitions are available where noted.'
)
doc.add_paragraph(
    'UPCOMING  →  SIGNUPS CLOSED  →  TEAMS PUBLISHED  →  COMPLETED\n'
    '     ↘                                                    \n'
    '   CANCELLED                                              '
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2 – WORKFLOW STAGE DETAIL
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Workflow Stages – Detailed Feature Matrix', level=1)

# ── 2.1 UPCOMING ────────────────────────────────────────────────────────────
doc.add_heading('2.1 Stage: UPCOMING (Sign-ups Open)', level=2)
doc.add_paragraph('The session has been created and is accepting RSVPs from players.')

doc.add_heading('Admin Actions', level=3)
add_table(
    ['#', 'Feature', 'Status', 'Description'],
    [
        ['1', 'Edit Session Details', 'EXISTS', 'Change date, venue, time, format, court cost, buffer %. Navigates to edit page.'],
        ['2', 'Add Player (Search)', 'EXISTS', 'Type-ahead search to find existing players and add to confirmed or waitlist.'],
        ['3', 'Add Player (Dropdown)', 'NEW', 'Dropdown list of all active players (alphabetical), filtered to those not in session. Alternative to search.'],
        ['4', 'Add New Player Inline', 'EXISTS', 'Create a brand-new player (name + mobile) and add to waitlist.'],
        ['5', 'Set Court Payer', 'EXISTS', 'Designate which player pays the court fee (auto-marks their payment as paid).'],
        ['6', 'Edit RSVP Status', 'NEW', 'Change a player\'s RSVP from confirmed → absent/maybe or vice versa without removing them.'],
        ['7', 'Remove Player from Session', 'EXISTS', 'Remove a player\'s RSVP entirely (they can re-RSVP).'],
        ['8', 'Promote Waitlist', 'EXISTS (improved)', 'Batch-promote waitlisted players to fill available spots. Now available during Upcoming (was only after close).'],
        ['9', 'Close Sign-ups', 'EXISTS', 'Transition to SIGNUPS CLOSED. Sets closed_at timestamp.'],
        ['10', 'Cancel Session', 'EXISTS', 'Transition to CANCELLED.'],
        ['11', 'Send Message', 'EXISTS', 'Message confirmed, waitlisted, or all active players via WhatsApp/SMS.'],
        ['12', 'Delete Session', 'EXISTS', 'Permanently remove session and all related data (RSVPs, payments, teams, notifications).'],
    ],
    col_widths=[1, 4, 2.5, 10]
)

doc.add_heading('Player Actions', level=3)
add_table(
    ['#', 'Feature', 'Status', 'Description'],
    [
        ['1', 'RSVP "I\'m In"', 'EXISTS', 'Join as confirmed. Auto-waitlisted if session is full.'],
        ['2', 'RSVP "Maybe"', 'EXISTS', 'Mark as maybe.'],
        ['3', 'RSVP "Can\'t Make It"', 'EXISTS', 'Mark as absent.'],
        ['4', 'Change RSVP', 'EXISTS', 'Switch between confirmed/maybe/absent.'],
        ['5', 'Withdraw', 'EXISTS', 'Remove self from session entirely.'],
        ['6', 'Add Other to Waitlist', 'EXISTS', 'Add another player (existing or new) to the waitlist.'],
        ['7', 'View Session Info', 'EXISTS', 'See venue, time, cost, confirmed list, waitlist.'],
    ],
    col_widths=[1, 4, 2.5, 10]
)

# ── 2.2 SIGNUPS CLOSED ──────────────────────────────────────────────────────
doc.add_heading('2.2 Stage: SIGNUPS CLOSED', level=2)
doc.add_paragraph('Sign-ups are frozen. Admin prepares for team generation.')

doc.add_heading('Admin Actions', level=3)
add_table(
    ['#', 'Feature', 'Status', 'Description'],
    [
        ['1', 'Edit Session Details', 'NEW', 'Allow editing venue, time, cost, buffer even after sign-ups close. If cost/buffer changes and payments exist, auto-recalculate all payment amounts.'],
        ['2', 'Add Player (Late Addition)', 'NEW', 'Admin can add a player directly to confirmed list, bypassing the waitlist and the max-player cap. For late arrivals or last-minute changes.'],
        ['3', 'Edit RSVP Status', 'NEW', 'Change a confirmed player to absent/maybe or vice versa.'],
        ['4', 'Remove Player from Session', 'NEW', 'Remove a confirmed player\'s RSVP (e.g., no-show notification received).'],
        ['5', 'Promote Waitlist', 'EXISTS', 'Batch-promote waitlisted players to fill available confirmed spots.'],
        ['6', 'Generate Teams', 'EXISTS', 'Run team balancing algorithm. Transitions to TEAMS PUBLISHED. Skips payment creation (already done at close).'],
        ['7', 'Reopen Sign-ups', 'NEW', 'Backward transition to UPCOMING. Clears closed_at. Deletes any unpaid payment records. Players can RSVP again.'],
        ['8', 'Remove All Maybes', 'NEW', 'Bulk action: remove all players with RSVP status "maybe" from the session.'],
        ['9', 'Send Reminder to Unpaid', 'NEW', 'Bulk action: send payment reminder to all unpaid players.'],
        ['10', 'Send Message', 'EXISTS', 'Message confirmed, waitlisted, unpaid, or all players.'],
        ['11', 'Delete Session', 'EXISTS', 'Permanently remove session.'],
    ],
    col_widths=[1, 4, 2.5, 10]
)

doc.add_heading('Player Actions', level=3)
add_table(
    ['#', 'Feature', 'Status', 'Description'],
    [
        ['1', 'View Session Info', 'EXISTS', 'See confirmed list, waitlist, venue, time.'],
        ['2', 'Pay via Card / PayID', 'NEW', 'Payments are auto-created when sign-ups close. Players can pay immediately.'],
    ],
    col_widths=[1, 4, 2.5, 10]
)

# ── 2.3 TEAMS PUBLISHED ─────────────────────────────────────────────────────
doc.add_heading('2.3 Stage: TEAMS PUBLISHED', level=2)
doc.add_paragraph('Teams have been generated and published. Players can see their team assignments and make payments.')

doc.add_heading('Admin Actions', level=3)
add_table(
    ['#', 'Feature', 'Status', 'Description'],
    [
        ['1', 'Edit Session Details', 'NEW', 'Edit venue, time, cost, buffer. If cost/buffer changes, auto-recalculate all unpaid payment amounts. Already-paid amounts are preserved.'],
        ['2', 'Regenerate Teams', 'EXISTS', 'Discard current teams and re-run the balancing algorithm.'],
        ['3', 'Swap Players Between Teams', 'EXISTS', 'Manually move a player from one team to another.'],
        ['4', 'Change Bib Colours', 'EXISTS', 'Assign bib colours to each team (White, Black, Red, Blue, Yellow, Green).'],
        ['5', 'Add Late Player to Team', 'NEW', 'Add a player directly to a specific team (admin picks the team). Creates their payment record automatically.'],
        ['6', 'Edit RSVP Status', 'NEW', 'Change a player\'s RSVP (e.g., mark a no-show as absent). Removes from team if changed to absent.'],
        ['7', 'Manage Payments (Stripe)', 'EXISTS', 'Process card payments with 3.5% surcharge via Stripe checkout.'],
        ['8', 'Manage Payments (PayID)', 'EXISTS', 'Player sends PayID transfer; admin approves/rejects pending confirmations.'],
        ['9', 'Manage Payments (Cash)', 'EXISTS', 'Admin marks individual payments as paid via cash.'],
        ['10', 'Chase Payments', 'EXISTS', 'Auto-escalation system: reminder at T+2d, escalation at T+5d, admin alert at T+7d.'],
        ['11', 'Pause/Resume Chase', 'EXISTS', 'Toggle auto-chase per player.'],
        ['12', 'Send Reminder to Unpaid', 'NEW', 'Bulk action: send payment reminder to all unpaid players in one click.'],
        ['13', 'Unpublish Teams', 'NEW', 'Backward transition to SIGNUPS CLOSED. Removes team assignments but preserves RSVPs and payments.'],
        ['14', 'Complete Session', 'EXISTS', 'Transition to COMPLETED. Shows warning if unpaid players remain.'],
        ['15', 'Send Message', 'EXISTS', 'Message groups or individuals.'],
        ['16', 'Delete Session', 'EXISTS', 'Permanently remove session.'],
    ],
    col_widths=[1, 4, 2.5, 10]
)

doc.add_heading('Player Actions', level=3)
add_table(
    ['#', 'Feature', 'Status', 'Description'],
    [
        ['1', 'View Team Assignment', 'EXISTS', 'See which team and teammates, bib colour.'],
        ['2', 'Pay via Card (Stripe)', 'EXISTS', 'Select unpaid players (self or dependants), pay with 3.5% card fee.'],
        ['3', 'Pay via PayID', 'EXISTS', 'See PayID number and reference code. Confirm payment, admin approves.'],
        ['4', 'View Payment Status', 'EXISTS', 'See paid/unpaid/pending status per player.'],
    ],
    col_widths=[1, 4, 2.5, 10]
)

# ── 2.4 COMPLETED ───────────────────────────────────────────────────────────
doc.add_heading('2.4 Stage: COMPLETED', level=2)
doc.add_paragraph('Session is finished. Read-only view with recap and historical data.')

doc.add_heading('Admin Actions', level=3)
add_table(
    ['#', 'Feature', 'Status', 'Description'],
    [
        ['1', 'View Session Recap', 'EXISTS', 'Attendance count, payment collection rate, amount collected vs due.'],
        ['2', 'Reopen Session', 'NEW', 'Backward transition to TEAMS PUBLISHED. Allows fixing payment errors or late payments.'],
        ['3', 'Send Reminder to Unpaid', 'NEW', 'Chase remaining unpaid players even after session completion.'],
        ['4', 'Delete Session', 'EXISTS', 'Permanently remove session.'],
    ],
    col_widths=[1, 4, 2.5, 10]
)

doc.add_heading('Player Actions', level=3)
add_table(
    ['#', 'Feature', 'Status', 'Description'],
    [
        ['1', 'View Session Recap', 'EXISTS', 'See final teams, attendance, payment summary.'],
        ['2', 'Pay Outstanding', 'NEW', 'If reopened, can still pay outstanding amounts.'],
    ],
    col_widths=[1, 4, 2.5, 10]
)

# ── 2.5 CANCELLED ───────────────────────────────────────────────────────────
doc.add_heading('2.5 Stage: CANCELLED', level=2)
doc.add_paragraph('Session has been cancelled. Minimal actions available.')

doc.add_heading('Admin Actions', level=3)
add_table(
    ['#', 'Feature', 'Status', 'Description'],
    [
        ['1', 'View Cancelled Session', 'EXISTS', 'See session details and who had RSVP\'d.'],
        ['2', 'Uncancel Session', 'NEW', 'Backward transition to UPCOMING. Restores RSVPs. Players can sign up again.'],
        ['3', 'Delete Session', 'EXISTS', 'Permanently remove session.'],
    ],
    col_widths=[1, 4, 2.5, 10]
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 3 – WORKFLOW TRANSITIONS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Workflow Transitions', level=1)

doc.add_heading('3.1 Forward Transitions', level=2)
add_table(
    ['From', 'To', 'Trigger', 'Side Effects'],
    [
        ['Upcoming', 'Signups Closed', 'Admin: "Close Sign-ups"', 'Sets closed_at timestamp. Players can no longer RSVP. Auto-creates payment records for all confirmed players (10 or 15). Court payer auto-marked as paid.'],
        ['Signups Closed', 'Teams Published', 'Admin: "Generate Teams"', 'Creates teams via balancing algorithm. Payment records already exist from sign-up close — no new payments created.'],
        ['Teams Published', 'Completed', 'Admin: "Complete Session"', 'Locks session. Warning shown if unpaid players exist.'],
        ['Upcoming', 'Cancelled', 'Admin: "Cancel"', 'Session marked cancelled. No data deleted.'],
    ],
    col_widths=[3, 3, 4, 7]
)

doc.add_heading('3.2 Backward Transitions (NEW)', level=2)
add_table(
    ['From', 'To', 'Trigger', 'Side Effects'],
    [
        ['Signups Closed', 'Upcoming', 'Admin: "Reopen Sign-ups"', 'Clears closed_at. Deletes unpaid payment records (paid ones preserved). Players can RSVP again. Existing RSVPs preserved.'],
        ['Teams Published', 'Signups Closed', 'Admin: "Unpublish Teams"', 'Removes team assignments (team records deleted). RSVPs preserved. Payment records preserved (amounts stay, can still be collected).'],
        ['Completed', 'Teams Published', 'Admin: "Reopen Session"', 'Unlocks session. Payments can be edited/collected again. Teams remain intact.'],
        ['Cancelled', 'Upcoming', 'Admin: "Uncancel"', 'Restores to Upcoming. All existing RSVPs preserved. Players can modify RSVPs.'],
    ],
    col_widths=[3, 3, 4, 7]
)

doc.add_heading('3.3 Transition Diagram', level=2)
doc.add_paragraph(
    '                    ┌──────────────┐\n'
    '          ┌────────→│   UPCOMING   │←───────────┐\n'
    '          │ Uncancel│  (Sign-ups)  │ Reopen     │\n'
    '          │         └──────┬───────┘ Sign-ups   │\n'
    '          │                │ Close                │\n'
    '          │                ↓                      │\n'
    '   ┌──────┴───┐    ┌──────────────┐              │\n'
    '   │CANCELLED │    │   SIGNUPS    │──────────────┘\n'
    '   └──────────┘    │   CLOSED     │\n'
    '                   └──────┬───────┘\n'
    '                          │ Generate Teams\n'
    '                  Unpub.  ↓         \n'
    '                ┌─────────────────┐\n'
    '                │TEAMS PUBLISHED  │←──────┐\n'
    '                └────────┬────────┘Reopen │\n'
    '                         │ Complete       │\n'
    '                         ↓                │\n'
    '                  ┌─────────────┐         │\n'
    '                  │  COMPLETED  │─────────┘\n'
    '                  └─────────────┘\n'
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 4 – PLAYER SEARCH & SELECTION
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Player Search & Selection Enhancement', level=1)

doc.add_heading('4.1 Current Behaviour', level=2)
doc.add_paragraph(
    'Players are added to a session via a type-ahead search field. The admin types a name '
    'and matching players appear in a filtered list below. This works well when the admin '
    'knows the player\'s name but is cumbersome when browsing.'
)

doc.add_heading('4.2 New: Dropdown Selector', level=2)
doc.add_paragraph('Add a dropdown (select list) alongside the existing search field:')
doc.add_paragraph('• Shows all active players not already in the session, sorted alphabetically.', style='List Bullet')
doc.add_paragraph('• Selecting a player from the dropdown adds them to the session (confirmed or waitlist based on capacity).', style='List Bullet')
doc.add_paragraph('• Both the search field and dropdown remain available — admin can use either method.', style='List Bullet')
doc.add_paragraph('• The dropdown updates in real-time as players are added/removed.', style='List Bullet')

doc.add_heading('4.3 Where This Applies', level=2)
add_table(
    ['Page', 'Section', 'Stages Available'],
    [
        ['Admin Session Detail', 'Add to Confirmed', 'Upcoming, Signups Closed, Teams Published'],
        ['Admin Session Detail', 'Add to Waitlist', 'Upcoming, Signups Closed'],
        ['Player Session Detail', 'Add Other to Waitlist', 'Upcoming'],
    ],
    col_widths=[4, 4, 6]
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 5 – RSVP MANAGEMENT
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('5. RSVP Management Enhancement', level=1)

doc.add_heading('5.1 Edit RSVP Status (NEW)', level=2)
doc.add_paragraph(
    'Admins should be able to change a player\'s RSVP status without removing and re-adding them. '
    'This is useful for marking no-shows, converting "maybe" responses, or correcting mistakes.'
)

add_table(
    ['Action', 'Available Stages', 'Side Effects'],
    [
        ['Confirmed → Absent', 'Upcoming, Closed, Teams Published', 'If in a team, player is removed from team. Opens a confirmed spot. Auto-promotes next waitlisted player if applicable.'],
        ['Confirmed → Maybe', 'Upcoming, Closed', 'Opens a confirmed spot. Auto-promotes next waitlisted player if applicable.'],
        ['Absent → Confirmed', 'Upcoming, Closed', 'Fills a confirmed spot (or goes to waitlist if full).'],
        ['Maybe → Confirmed', 'Upcoming, Closed', 'Fills a confirmed spot (or goes to waitlist if full).'],
        ['Maybe → Absent', 'Upcoming, Closed', 'No capacity change.'],
    ],
    col_widths=[4, 4, 9]
)

doc.add_heading('5.2 Remove All Maybes (NEW)', level=2)
doc.add_paragraph(
    'Bulk action available during Upcoming and Signups Closed stages. '
    'Removes all RSVPs with status "maybe" from the session in one click. '
    'Confirmation dialog shown before executing: "Remove X player(s) with \'maybe\' status?"'
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 6 – PAYMENT ENHANCEMENTS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Payment Enhancements', level=1)

doc.add_heading('6.1 Auto-Create Payments on Sign-up Close (NEW)', level=2)
doc.add_paragraph(
    'Currently, payment records are only created when teams are published. '
    'With this change, payments are automatically created when the admin closes sign-ups. '
    'By this point, the roster is finalized at the max capacity (10 for 2 teams, 15 for 3 teams), '
    'so the per-player cost is known.'
)
doc.add_paragraph('• Closing sign-ups automatically creates payment records for all confirmed players.', style='List Bullet')
doc.add_paragraph('• Payment amount = court_cost × (1 + buffer_pct/100) ÷ max_players.', style='List Bullet')
doc.add_paragraph('• Court payer auto-marked as paid.', style='List Bullet')
doc.add_paragraph('• Players can start paying immediately via Stripe or PayID — no need to wait for team generation.', style='List Bullet')
doc.add_paragraph('• Generate Teams no longer creates payments (already done).', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Reopening sign-ups (backward transition) deletes all unpaid payment records. '
    'Paid payments are preserved. When sign-ups are closed again, new payment records are created '
    'for confirmed players who don\'t already have a payment record.')

doc.add_heading('6.2 Auto-Recalculate on Cost Change (NEW)', level=2)
doc.add_paragraph(
    'If the admin edits the court cost or buffer percentage after payments have been created:'
)
doc.add_paragraph('• All UNPAID payment amounts are recalculated with the new cost.', style='List Bullet')
doc.add_paragraph('• Already PAID payments are NOT changed (amount_paid is preserved).', style='List Bullet')
doc.add_paragraph('• A toast notification confirms: "X payment(s) recalculated to $Y.YY each."', style='List Bullet')
doc.add_paragraph('• Pending PayID confirmations remain pending but with updated amount_due.', style='List Bullet')

doc.add_heading('6.3 Bulk Reminder to Unpaid (NEW)', level=2)
doc.add_paragraph(
    'A "Remind All Unpaid" button available during Signups Closed, Teams Published, and Completed stages. '
    'Sends a payment reminder via WhatsApp/SMS to all players with unpaid status (excluding court payer '
    'and those with chase_paused = true).'
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 7 – LATE PLAYER ADDITIONS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Late Player Additions', level=1)

doc.add_heading('7.1 Add After Sign-ups Closed (NEW)', level=2)
doc.add_paragraph(
    'Admin can add a player directly to the confirmed list even after sign-ups are closed. '
    'This bypasses the max-player cap and waitlist. Use case: a regular player confirms late.'
)

doc.add_heading('7.2 Add After Teams Published (NEW)', level=2)
doc.add_paragraph(
    'Admin can add a late player and assign them to a specific team:'
)
doc.add_paragraph('1. Admin searches/selects the player (using search or dropdown).', style='List Number')
doc.add_paragraph('2. A team picker appears: "Assign to Team A / Team B / Team C".', style='List Number')
doc.add_paragraph('3. Player is added to the selected team.', style='List Number')
doc.add_paragraph('4. A payment record is automatically created for the new player.', style='List Number')
doc.add_paragraph('5. Team average skill ratings are recalculated.', style='List Number')

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 8 – SESSION EDITING
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Session Editing at Later Stages', level=1)

doc.add_heading('8.1 Current Behaviour', level=2)
doc.add_paragraph(
    'The "Edit" button and edit page are only available when session status is "upcoming". '
    'Once sign-ups close, session details cannot be changed.'
)

doc.add_heading('8.2 New: Edit at All Active Stages', level=2)
doc.add_paragraph('The Edit button will be available during Upcoming, Signups Closed, and Teams Published stages.')

add_table(
    ['Field', 'Editable Stages', 'Notes'],
    [
        ['Date', 'Upcoming, Closed, Teams Published', 'No side effects on existing data.'],
        ['Venue', 'Upcoming, Closed, Teams Published', 'No side effects.'],
        ['Start / End Time', 'Upcoming, Closed, Teams Published', 'No side effects.'],
        ['Format (2t / 3t)', 'Upcoming, Closed', 'NOT editable after teams published (would invalidate teams). Changing format updates max_players.'],
        ['Court Cost', 'Upcoming, Closed, Teams Published', 'If payments exist, triggers recalculation of unpaid amounts.'],
        ['Buffer %', 'Upcoming, Closed, Teams Published', 'If payments exist, triggers recalculation of unpaid amounts.'],
    ],
    col_widths=[3, 5, 9]
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 9 – FEATURE SUMMARY BY STAGE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Complete Feature Matrix', level=1)
doc.add_paragraph('Summary of all features and their availability across workflow stages.')
doc.add_paragraph('Key: E = Exists, N = New, — = Not Available')

add_table(
    ['Feature', 'Upcoming', 'Closed', 'Teams Pub.', 'Completed', 'Cancelled'],
    [
        ['Edit Session Details', 'E', 'N', 'N', '—', '—'],
        ['Add Player (Search)', 'E', 'N', 'N', '—', '—'],
        ['Add Player (Dropdown)', 'N', 'N', 'N', '—', '—'],
        ['Edit RSVP Status', 'N', 'N', 'N', '—', '—'],
        ['Remove Player', 'E', 'N', 'N', '—', '—'],
        ['Promote Waitlist (Batch)', 'E', 'E', '—', '—', '—'],
        ['Remove All Maybes', 'N', 'N', '—', '—', '—'],
        ['Auto-Create Payments', '—', 'N (on close)', 'E (exist)', '—', '—'],
        ['Pay (Card / PayID)', '—', 'N', 'E', 'N', '—'],
        ['Pay (Cash - Admin)', '—', 'N', 'E', 'N', '—'],
        ['Remind All Unpaid', '—', 'N', 'N', 'N', '—'],
        ['Chase Payments (Auto)', '—', '—', 'E', '—', '—'],
        ['Generate Teams', '—', 'E', '—', '—', '—'],
        ['Regenerate Teams', '—', '—', 'E', '—', '—'],
        ['Swap Players / Bibs', '—', '—', 'E', '—', '—'],
        ['Add Late Player to Team', '—', '—', 'N', '—', '—'],
        ['Complete Session', '—', '—', 'E', '—', '—'],
        ['View Recap', '—', '—', '—', 'E', '—'],
        ['Send Messages', 'E', 'E', 'E', 'N', '—'],
        ['Close Sign-ups', 'E', '—', '—', '—', '—'],
        ['Cancel Session', 'E', '—', '—', '—', '—'],
        ['Reopen Sign-ups', '—', 'N', '—', '—', '—'],
        ['Unpublish Teams', '—', '—', 'N', '—', '—'],
        ['Reopen Session', '—', '—', '—', 'N', '—'],
        ['Uncancel', '—', '—', '—', '—', 'N'],
        ['Delete Session', 'E', 'E', 'E', 'E', 'E'],
        ['Set Court Payer', 'E', 'E', 'E', '—', '—'],
    ],
    col_widths=[4.5, 2, 2, 2, 2, 2]
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 10 – UI BUTTON PATTERNS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('10. UI Button Visibility & Disabled Patterns', level=1)
doc.add_paragraph(
    'A consistent two-tier approach is used across the app to determine whether buttons are '
    'hidden entirely or shown in a disabled (greyed-out) state.'
)

doc.add_heading('10.1 Pattern Rules', level=2)
add_table(
    ['Scenario', 'Approach', 'Reasoning'],
    [
        ['Wrong workflow stage', 'Disabled (greyed) + tooltip', 'User sees what actions exist and understands why they are not available yet. E.g., "Generate Teams" visible but greyed with tooltip "Close sign-ups first".'],
        ['Loading / processing', 'Disabled (greyed)', 'Prevents double-clicks. Already done consistently.'],
        ['No permission (non-admin)', 'Hidden entirely', 'Non-admins should not see admin-only features.'],
        ['No data (empty waitlist, no pending PayID)', 'Hidden entirely', 'No point showing an action when there is nothing to act on.'],
        ['Already actioned (current RSVP)', 'Disabled (greyed)', 'User sees all options but knows which is currently selected.'],
    ],
    col_widths=[4, 4, 9]
)

doc.add_heading('10.2 Workflow Button Visibility by Stage', level=2)
doc.add_paragraph(
    'All workflow action buttons are always rendered for admins. Buttons that cannot be used '
    'at the current stage are greyed out with a tooltip explaining the prerequisite.'
)
add_table(
    ['Button', 'Upcoming', 'Closed', 'Teams Pub.', 'Completed', 'Cancelled'],
    [
        ['Edit Session', 'Enabled', 'Enabled', 'Enabled', 'Disabled', 'Disabled'],
        ['Close Sign-ups', 'Enabled', 'Disabled (already closed)', 'Disabled', 'Disabled', 'Disabled'],
        ['Reopen Sign-ups', 'Disabled (already open)', 'Enabled', 'Disabled', 'Disabled', 'Disabled'],
        ['Generate Teams', 'Disabled (close sign-ups first)', 'Enabled', 'Disabled (already published)', 'Disabled', 'Disabled'],
        ['Unpublish Teams', 'Disabled', 'Disabled', 'Enabled', 'Disabled', 'Disabled'],
        ['Complete Session', 'Disabled', 'Disabled', 'Enabled', 'Disabled (already done)', 'Disabled'],
        ['Reopen Session', 'Disabled', 'Disabled', 'Disabled', 'Enabled', 'Disabled'],
        ['Cancel Session', 'Enabled', 'Disabled', 'Disabled', 'Disabled', 'Disabled'],
        ['Uncancel', 'Disabled', 'Disabled', 'Disabled', 'Disabled', 'Enabled'],
        ['Delete Session', 'Enabled', 'Enabled', 'Enabled', 'Enabled', 'Enabled'],
    ],
    col_widths=[3.5, 2.3, 2.3, 2.3, 2.3, 2.3]
)

doc.add_heading('10.3 Tooltip Messages', level=2)
add_table(
    ['Button (when disabled)', 'Tooltip Text'],
    [
        ['Edit Session', '"Session is closed for editing"'],
        ['Close Sign-ups', '"Sign-ups already closed"'],
        ['Reopen Sign-ups', '"Sign-ups are already open"'],
        ['Generate Teams', '"Close sign-ups first"'],
        ['Unpublish Teams', '"No teams to unpublish"'],
        ['Complete Session', '"Publish teams first"'],
        ['Reopen Session', '"Session is not completed"'],
        ['Cancel Session', '"Only upcoming sessions can be cancelled"'],
        ['Uncancel', '"Session is not cancelled"'],
    ],
    col_widths=[5, 12]
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 11 – IMPLEMENTATION PRIORITY
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('11. Implementation Priority', level=1)
doc.add_paragraph(
    'Recommended order of implementation, grouped by effort and impact.'
)

doc.add_heading('11.1 Quick Wins (Small effort, high value)', level=2)
doc.add_paragraph('1. Player dropdown selector (alongside existing search)', style='List Number')
doc.add_paragraph('2. Promote Waitlist button during Upcoming stage (DONE)', style='List Number')
doc.add_paragraph('3. Remove All Maybes bulk action', style='List Number')

doc.add_heading('11.2 Medium Effort', level=2)
doc.add_paragraph('4. Edit RSVP status (admin changes confirmed / absent / maybe)', style='List Number')
doc.add_paragraph('5. Auto-create payments when sign-ups close', style='List Number')
doc.add_paragraph('6. Edit session details at later stages (with payment recalculation)', style='List Number')
doc.add_paragraph('7. Late player addition after sign-ups closed (direct to confirmed)', style='List Number')

doc.add_heading('11.3 Larger Features', level=2)
doc.add_paragraph('8. Backward transitions (reopen sign-ups, unpublish teams, reopen completed, uncancel)', style='List Number')
doc.add_paragraph('9. Add late player to specific team (with payment auto-creation)', style='List Number')
doc.add_paragraph('10. Consistent button disabled/tooltip pattern across all workflow stages', style='List Number')

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 12 – DECISIONS LOG
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('12. Decisions Log', level=1)
doc.add_paragraph('Key decisions made during requirements gathering:')

add_table(
    ['#', 'Decision', 'Rationale'],
    [
        ['1', 'All four backward transitions approved', 'Flexibility to correct mistakes at any stage.'],
        ['2', 'RSVP edit = change status only, not player profile', 'Player profile edits belong on the Players page, not mid-session.'],
        ['3', 'Late additions bypass max-player cap', 'Admin judgment call — real-world flexibility.'],
        ['4', 'Late additions to teams: admin picks the team', 'No auto-assignment; admin controls team balance.'],
        ['5', 'Payments auto-created when sign-ups close', 'Roster is finalized at max capacity (10 or 15) by close time, so per-player cost is known. No reason to wait for teams.'],
        ['6', 'Session edit available through Teams Published', 'Venue/time/cost can change up until completion.'],
        ['7', 'Format not editable after teams published', 'Would invalidate existing team structure.'],
        ['8', 'Cost changes auto-recalculate unpaid only', 'Paid amounts are committed and not retroactively changed.'],
        ['9', 'No match results or man-of-the-match for now', 'Future consideration — not needed in current phase.'],
        ['10', 'Dropdown shows all active players alphabetically', 'Simple and predictable. Can enhance later with frequency sorting.'],
        ['11', 'Buttons: disabled+tooltip for wrong stage, hidden for no permission/no data', 'Admins see the full lifecycle of actions. Greyed buttons with tooltips teach the workflow and reduce "where did that button go?" confusion.'],
        ['12', 'Removed "Remind All Unpaid" bulk action', 'Not practical — everyone starts unpaid after close. The automated chase system (T+2, T+5, T+7) already handles reminders.'],
    ],
    col_widths=[1, 6, 10]
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 13 – APPENDIX
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('13. Appendix', level=1)

doc.add_heading('13.1 Session Statuses', level=2)
add_table(
    ['Status', 'DB Value', 'Description'],
    [
        ['Upcoming', 'upcoming', 'Session created, accepting RSVPs'],
        ['Signups Closed', 'signups_closed', 'RSVPs frozen, preparing for teams'],
        ['Teams Published', 'teams_published', 'Teams generated, payments active'],
        ['Completed', 'completed', 'Session finished, read-only'],
        ['Cancelled', 'cancelled', 'Session cancelled'],
    ],
    col_widths=[3, 4, 10]
)

doc.add_heading('13.2 Payment Statuses', level=2)
add_table(
    ['Status', 'DB Value', 'Description'],
    [
        ['Unpaid', 'unpaid', 'Payment not yet made'],
        ['Paid', 'paid', 'Payment received and confirmed'],
        ['Pending Confirmation', 'pending_confirmation', 'PayID transfer claimed, awaiting admin approval'],
    ],
    col_widths=[4, 4, 9]
)

doc.add_heading('13.3 Payment Methods', level=2)
add_table(
    ['Method', 'DB Value', 'Fee', 'Description'],
    [
        ['Stripe (Card)', 'stripe', '3.5% surcharge', 'Online card payment via Stripe checkout'],
        ['PayID', 'payid', 'None', 'Bank transfer using PayID phone number + reference code'],
        ['Cash', 'cash', 'None', 'In-person cash payment, admin marks as paid'],
    ],
    col_widths=[3, 2.5, 3, 8.5]
)

doc.add_heading('13.4 RSVP Statuses', level=2)
add_table(
    ['Status', 'DB Value', 'Description'],
    [
        ['Confirmed', 'confirmed', 'Player is attending (or waitlisted if is_waitlist=true)'],
        ['Maybe', 'maybe', 'Player is uncertain'],
        ['Absent', 'absent', 'Player cannot attend'],
    ],
    col_widths=[3, 4, 10]
)

doc.add_heading('13.5 Player Types', level=2)
add_table(
    ['Type', 'DB Value', 'UI Colour', 'Description'],
    [
        ['Regular', 'regular', 'Blue', 'Core group member, plays most weeks'],
        ['Casual', 'casual', 'Orange', 'Occasional or one-off player'],
        ['Admin', '(is_admin flag)', 'Purple', 'Has admin privileges, can manage sessions'],
    ],
    col_widths=[3, 3, 2.5, 8.5]
)

# ── Save ────────────────────────────────────────────────────────────────────
output_path = r'C:\Users\MagedBoctor\Claude\Soccer App\docs\Monday_Night_Soccer_Requirements.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
