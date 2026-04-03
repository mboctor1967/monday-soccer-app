#!/usr/bin/env python3
"""Generate the Monday Night Soccer App – Test Cases document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def add_test_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

HEADERS = ['#', 'Test Case', 'Player Type', 'Method', 'Status', 'Result', 'Notes']
COL_W = [1, 6, 2, 1.5, 1.5, 1.5, 3]

# ── Title Page ──────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Monday Night Soccer App')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Test Cases & Results')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Version 1.0  •  {datetime.date.today().strftime("%d %B %Y")}').font.size = Pt(12)

doc.add_paragraph()
legend = doc.add_paragraph()
legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
legend.add_run('Status: ').bold = True
legend.add_run('PASS / FAIL / SKIP / BLOCKED / NOT RUN')
doc.add_paragraph()
method_legend = doc.add_paragraph()
method_legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
method_legend.add_run('Method: ').bold = True
method_legend.add_run('Auto = automated test  |  Manual = manual UI testing')

doc.add_page_break()

# ── Test Data ───────────────────────────────────────────────────────────────
doc.add_heading('1. Test Data — Player Roster', level=1)
doc.add_paragraph('The following players exist in the test database and are used across test cases.')

doc.add_heading('1.1 Admin Players', level=2)
add_test_table(
    ['Name', 'Mobile', 'Email', 'Type', 'Skill', 'Active'],
    [
        ['David Ibrahim', '+61449950350', 'david_ibrahim7@hotmail.com', 'Regular', '4', 'Yes'],
        ['Karim Azer', '+61418612568', 'karimazer3@gmail.com', 'Regular', '3', 'Yes'],
        ['Maged Boctor', '+61412408587', 'mboctor@gmail.com', 'Regular', '3', 'Yes'],
        ['Marcus Boctor', '+61444517626', 'marcusboctor03@gmail.com', 'Regular', '4', 'Yes'],
        ['Michael Boctor', '+61452508587', 'michael.boctor@gmail.com', 'Regular', '5', 'Yes'],
        ['Remon Metira', '+61404223113', 'remon.metira@gmail.com', 'Regular', '3', 'Yes'],
    ],
    col_widths=[3.5, 3, 5, 1.8, 1, 1]
)

doc.add_heading('1.2 Non-Admin Regular Players', level=2)
add_test_table(
    ['Name', 'Mobile', 'Type', 'Skill', 'Active'],
    [
        ['Andrew Beshara', '+61490333825', 'Regular', '3', 'Yes'],
        ['Andy Makram', '+61452505861', 'Regular', '3', 'Yes'],
        ['Daniel Ishak', '+61431532531', 'Regular', '3', 'Yes'],
        ['Daniel Kilada', '+61450736341', 'Regular', '3', 'Yes'],
        ['David Mikhail', '+61416016764', 'Regular', '3', 'Yes'],
        ['Fady Girgis', '+61403246429', 'Regular', '3', 'Yes'],
        ['George Zikry', '+971507276467', 'Regular', '3', 'Yes'],
        ['Isaac Nayrouz', '+61423170023', 'Regular', '3', 'Yes'],
        ['Jeremy Nada', '+61401969732', 'Regular', '3', 'Yes'],
        ['Jonathan Tanios', '+61416930490', 'Regular', '3', 'Yes'],
        ['Marc Bastaworous', '+61420309091', 'Regular', '3', 'Yes'],
        ['Phillo Mikhail', '+61403352922', 'Regular', '3', 'Yes'],
    ],
    col_widths=[3.5, 3, 2, 1.5, 1.5]
)

doc.add_heading('1.3 Casual Players', level=2)
add_test_table(
    ['Name', 'Mobile', 'Type', 'Skill', 'Active'],
    [
        ['Billy Michael_Friend', '+61490187274', 'Casual', '3', 'Yes'],
    ],
    col_widths=[3.5, 3, 2, 1.5, 1.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION A – Authentication
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Test Cases', level=1)

doc.add_heading('A. Authentication & Login', level=2)
add_test_table(HEADERS, [
    ['A1', 'Admin login with OTP bypass (DEV_SKIP_OTP=true) — enter any code, should authenticate', 'Admin', 'Auto', '', '', ''],
    ['A2', 'Non-admin regular player login — enter mobile, bypass OTP, should see player dashboard', 'Regular', 'Auto', '', '', ''],
    ['A3', 'Casual player login — enter mobile, bypass OTP, should see player dashboard', 'Casual', 'Auto', '', '', ''],
    ['A4', 'Invalid mobile number — should show error, not create session', 'Any', 'Auto', '', '', ''],
], COL_W)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION B – Session Viewing
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('B. Session Viewing (Player)', level=2)
add_test_table(HEADERS, [
    ['B1', 'Regular player sees active sessions list with correct dates and venues', 'Regular', 'Auto', '', '', ''],
    ['B2', 'Casual player sees active sessions list', 'Casual', 'Auto', '', '', ''],
    ['B3', 'Player sees session detail page: venue, time, cost per player', 'Any', 'Auto', '', '', ''],
    ['B4', 'Admin player sees "Admin" button on session detail page', 'Admin', 'Manual', '', '', ''],
    ['B5', 'Non-admin player does NOT see "Admin" button', 'Regular', 'Manual', '', '', ''],
], COL_W)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION C – RSVP Flow
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('C. RSVP Flow (Player)', level=2)
add_test_table(HEADERS, [
    ['C1', 'Player RSVPs "I\'m In" — added to confirmed list, count increases', 'Regular', 'Auto', '', '', ''],
    ['C2', 'Player RSVPs when session is full (10/10 or 15/15) — auto-waitlisted with correct position', 'Regular', 'Auto', '', '', ''],
    ['C3', 'Player changes RSVP to "Maybe" — moved from confirmed to maybe section', 'Regular', 'Auto', '', '', ''],
    ['C4', 'Player changes RSVP to "Can\'t Make It" — moved to absent section', 'Regular', 'Auto', '', '', ''],
    ['C5', 'Player withdraws from session — RSVP deleted, no longer in any list', 'Regular', 'Auto', '', '', ''],
    ['C6', 'Casual player RSVPs — row shows orange background in UI', 'Casual', 'Manual', '', '', ''],
    ['C7', 'Player cannot RSVP after sign-ups closed — RSVP buttons not shown', 'Regular', 'Manual', '', '', ''],
], COL_W)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION D – Session Management
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('D. Admin — Session Management', level=2)
add_test_table(HEADERS, [
    ['D1', 'Create new session with default values — date, venue, format, cost, buffer', 'Admin', 'Auto', '', '', ''],
    ['D2', 'Edit session during Upcoming — change venue, time, cost. Verify saved correctly', 'Admin', 'Auto', '', '', ''],
    ['D3', 'Edit session during Signups Closed — change court cost. Verify unpaid payments recalculated, paid unchanged', 'Admin', 'Auto', '', '', ''],
    ['D4', 'Edit session during Teams Published — format dropdown is disabled/locked', 'Admin', 'Manual', '', '', ''],
    ['D5', 'Close sign-ups — status changes to signups_closed, payment records created for all confirmed players', 'Admin', 'Auto', '', '', ''],
    ['D6', 'Cancel session — status changes to cancelled', 'Admin', 'Auto', '', '', ''],
    ['D7', 'Delete session — session and all related data (RSVPs, payments, teams) deleted', 'Admin', 'Auto', '', '', ''],
], COL_W)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION E – Player Management in Session
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('E. Admin — Player Management in Session', level=2)
add_test_table(HEADERS, [
    ['E1', 'Add player via dropdown — alphabetical list shows only players not in session', 'Admin', 'Manual', '', '', ''],
    ['E2', 'Add player via type-ahead search — filtered results, click to add', 'Admin', 'Manual', '', '', ''],
    ['E3', 'Add player to waitlist via dropdown — goes to waitlist with correct position', 'Admin', 'Manual', '', '', ''],
    ['E4', 'Remove confirmed player — RSVP deleted, player no longer in session', 'Admin', 'Auto', '', '', ''],
    ['E5', 'Change RSVP: confirmed → absent. If on team, removed from team. Waitlist auto-promoted', 'Admin', 'Auto', '', '', ''],
    ['E6', 'Change RSVP: absent → confirmed. Added to confirmed if spots available', 'Admin', 'Auto', '', '', ''],
    ['E7', 'Change RSVP: maybe → confirmed when session full — goes to waitlist instead', 'Admin', 'Auto', '', '', ''],
    ['E8', 'Remove All Maybes — confirmation dialog shown, all "maybe" RSVPs deleted', 'Admin', 'Manual', '', '', ''],
    ['E9', 'Promote Waitlist during Upcoming — batch promotes up to available spots', 'Admin', 'Auto', '', '', ''],
    ['E10', 'Promote Waitlist during Signups Closed — same behaviour', 'Admin', 'Auto', '', '', ''],
    ['E11', 'Late player addition after sign-ups closed — bypasses max cap, payment auto-created', 'Admin', 'Auto', '', '', ''],
    ['E12', 'Late player added to specific team (teams published) — team_player created, team avg recalculated', 'Admin', 'Auto', '', '', ''],
    ['E13', 'Set court payer — payment status auto-set to "paid", amount_paid = amount_due', 'Admin', 'Auto', '', '', ''],
], COL_W)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION F – Backward Transitions
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('F. Admin — Backward Transitions', level=2)
add_test_table(HEADERS, [
    ['F1', 'Reopen Sign-ups (signups_closed → upcoming): unpaid payments deleted, paid preserved, closed_at cleared', 'Admin', 'Auto', '', '', ''],
    ['F2', 'Unpublish Teams (teams_published → signups_closed): teams and team_players deleted, RSVPs and payments preserved', 'Admin', 'Auto', '', '', ''],
    ['F3', 'Reopen Session (completed → teams_published): status changes, session unlocked for edits', 'Admin', 'Auto', '', '', ''],
    ['F4', 'Uncancel (cancelled → upcoming): status restored, existing RSVPs preserved', 'Admin', 'Auto', '', '', ''],
    ['F5', 'Round-trip: close → reopen → close again. Payments deleted on reopen, recreated on second close', 'Admin', 'Auto', '', '', ''],
], COL_W)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION G – Team Generation
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('G. Team Generation', level=2)
add_test_table(HEADERS, [
    ['G1', 'Generate 2 teams with 10 confirmed players — 5 per team, teams created in DB', 'Admin', 'Auto', '', '', ''],
    ['G2', 'Generate 3 teams with 15 confirmed players — 5 per team', 'Admin', 'Auto', '', '', ''],
    ['G3', 'Teams are skill-balanced — max team avg minus min team avg is within tolerance', 'Admin', 'Auto', '', '', ''],
    ['G4', 'Reshuffle teams — new proposals generated with different player distribution', 'Admin', 'Manual', '', '', ''],
    ['G5', 'Swap player between teams — drag/click swap updates team_players', 'Admin', 'Manual', '', '', ''],
    ['G6', 'Change bib colours — team record updated with new colour', 'Admin', 'Manual', '', '', ''],
    ['G7', 'Publish teams — status transitions to teams_published, teams saved to DB', 'Admin', 'Auto', '', '', ''],
], COL_W)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION H – Payments
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('H. Payments', level=2)
add_test_table(HEADERS, [
    ['H1', 'Payments auto-created on close — correct count, correct amount (court_cost * (1+buffer%) / max_players)', 'Admin', 'Auto', '', '', ''],
    ['H2', 'Court payer payment auto-marked as paid with amount_paid = amount_due', 'Admin', 'Auto', '', '', ''],
    ['H3', 'Pay via Stripe — redirects to Stripe checkout, returns with payment=success, verify-checkout marks as paid', 'Regular', 'Manual', '', '', ''],
    ['H4', 'Pay via PayID — status becomes pending_confirmation, admin approves → status becomes paid', 'Regular+Admin', 'Manual', '', '', ''],
    ['H5', 'Pay via PayID — admin rejects → status returns to unpaid, player can retry', 'Admin', 'Auto', '', '', ''],
    ['H6', 'Admin marks Cash payment — payment_status set to paid, payment_method set to cash', 'Admin', 'Auto', '', '', ''],
    ['H7', 'Cost change recalculates unpaid payments only — paid amounts preserved, unpaid get new amount_due', 'Admin', 'Auto', '', '', ''],
    ['H8', 'Player sees correct paid/unpaid/pending badges on session detail page', 'Any', 'Manual', '', '', ''],
    ['H9', 'Stripe webhook (checkout.session.completed) marks payment as paid in production', 'System', 'Manual', '', '', 'Requires Stripe CLI'],
    ['H10', 'Stripe verify-checkout endpoint confirms payment and updates DB (local fallback)', 'System', 'Auto', '', '', ''],
], COL_W)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION I – Messaging
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('I. Messaging', level=2)
add_test_table(HEADERS, [
    ['I1', 'Send message to all confirmed players — notification records created for each player', 'Admin', 'Auto', '', '', ''],
    ['I2', 'Send message to waitlisted only — only waitlisted players receive notification', 'Admin', 'Auto', '', '', ''],
    ['I3', 'Send message to unpaid only — only unpaid (non-court-payer) players targeted', 'Admin', 'Auto', '', '', ''],
    ['I4', 'Copy message to clipboard — clipboard contains expected text', 'Admin', 'Manual', '', '', ''],
], COL_W)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION J – Profile & Stats
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('J. Profile & Stats', level=2)
add_test_table(HEADERS, [
    ['J1', 'Player views own profile — shows name, mobile, email, attendance rate, payment stats', 'Regular', 'Manual', '', '', ''],
    ['J2', 'Player edits own name and email — changes saved to DB', 'Regular', 'Auto', '', '', ''],
    ['J3', 'Admin sees skill ratings (stars) on player rows in session detail', 'Admin', 'Manual', '', '', ''],
], COL_W)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 3 – Summary
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Test Execution Summary', level=1)

doc.add_heading('3.1 Results by Category', level=2)
add_test_table(
    ['Category', 'Total', 'Auto', 'Manual', 'Pass', 'Fail', 'Skip', 'Not Run'],
    [
        ['A. Authentication', '4', '3', '1', '', '', '', '4'],
        ['B. Session Viewing', '5', '2', '3', '', '', '', '5'],
        ['C. RSVP Flow', '7', '5', '2', '', '', '', '7'],
        ['D. Session Management', '7', '5', '2', '', '', '', '7'],
        ['E. Player Mgmt in Session', '13', '9', '4', '', '', '', '13'],
        ['F. Backward Transitions', '5', '5', '0', '', '', '', '5'],
        ['G. Team Generation', '7', '3', '4', '', '', '', '7'],
        ['H. Payments', '10', '6', '4', '', '', '', '10'],
        ['I. Messaging', '4', '3', '1', '', '', '', '4'],
        ['J. Profile & Stats', '3', '1', '2', '', '', '', '3'],
        ['TOTAL', '65', '42', '23', '', '', '', '65'],
    ],
    col_widths=[4, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5]
)

doc.add_heading('3.2 Overall Status', level=2)
add_test_table(
    ['Metric', 'Value'],
    [
        ['Total Test Cases', '65'],
        ['Automated', '42 (65%)'],
        ['Manual', '23 (35%)'],
        ['Pass', ''],
        ['Fail', ''],
        ['Skip', ''],
        ['Not Run', '65'],
        ['Pass Rate', ''],
        ['Test Run Date', ''],
        ['Tester', ''],
        ['Environment', 'localhost:3002 / Supabase (test)'],
    ],
    col_widths=[5, 10]
)

doc.add_heading('3.3 Defects Found', level=2)
add_test_table(
    ['#', 'Test Case', 'Severity', 'Description', 'Status', 'Fix Date'],
    [
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
    ],
    col_widths=[1, 2, 2, 6, 2, 2]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 4 – Test Environment
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Test Environment', level=1)
add_test_table(
    ['Component', 'Detail'],
    [
        ['App URL (Local)', 'http://localhost:3002'],
        ['App URL (Production)', 'Vercel deployment'],
        ['Database', 'Supabase (punekqewhvthvzmjwsvr)'],
        ['Stripe Mode', 'Test mode (sk_test_*, pk_test_*)'],
        ['OTP', 'Bypassed (DEV_SKIP_OTP=true)'],
        ['Test Framework', 'Vitest (API-level integration tests)'],
        ['Browser Testing', 'Manual'],
    ],
    col_widths=[4, 12]
)

doc.add_heading('4.1 Test Card Numbers (Stripe Test Mode)', level=2)
add_test_table(
    ['Card', 'Number', 'Result'],
    [
        ['Visa (success)', '4242 4242 4242 4242', 'Payment succeeds'],
        ['Visa (decline)', '4000 0000 0000 0002', 'Payment declined'],
        ['3D Secure', '4000 0025 0000 3155', 'Requires authentication'],
        ['Insufficient funds', '4000 0000 0000 9995', 'Declined — insufficient funds'],
    ],
    col_widths=[4, 5, 7]
)
doc.add_paragraph('Use any future expiry date (e.g. 12/34), any 3-digit CVC, and any postcode.')

# ── Save ────────────────────────────────────────────────────────────────────
output_path = r'C:\Users\MagedBoctor\Claude\Soccer App\docs\Monday_Night_Soccer_Test_Cases.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
